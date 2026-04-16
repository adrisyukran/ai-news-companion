"""
Test suite for ParserService.

Tests cover:
- Plain text parsing
- URL extraction
- PDF parsing
- DOCX parsing
- Error handling
"""

import pytest
from pathlib import Path
from backend.services.parser_service import ParserService, ParserError


class TestPlainTextParsing:
    """Tests for plain text handling."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
    
    def test_parse_simple_text(self):
        """Test parsing simple plain text."""
        text = "This is a sample article."
        result = self.service.parse(text)
        assert result == text
    
    def test_parse_text_with_whitespace(self):
        """Test that leading/trailing whitespace is stripped."""
        text = "  \n  This is text with whitespace.  \n  "
        result = self.service.parse(text)
        assert result == "This is text with whitespace."
    
    def test_parse_empty_text_raises_error(self):
        """Test that empty text raises ParserError."""
        with pytest.raises(ParserError):
            self.service.parse("")
    
    def test_parse_whitespace_only_raises_error(self):
        """Test that whitespace-only text raises ParserError."""
        with pytest.raises(ParserError):
            self.service.parse("   \n\t  ")
    
    def test_parse_multiline_text(self):
        """Test parsing multiline plain text."""
        text = """Line 1
Line 2
Line 3"""
        result = self.service.parse(text)
        assert "Line 1" in result
        assert "Line 2" in result
        assert "Line 3" in result


class TestURLParsing:
    """Tests for URL extraction."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
    
    def test_is_url_http(self):
        """Test HTTP URL detection."""
        assert self.service._is_url("http://example.com") is True
    
    def test_is_url_https(self):
        """Test HTTPS URL detection."""
        assert self.service._is_url("https://example.com/article") is True
    
    def test_is_url_not_url(self):
        """Test non-URL detection."""
        assert self.service._is_url("not a url") is False
        assert self.service._is_url("ftp://example.com") is False
    
    def test_parse_url_connection_error_raises_error(self):
        """Test that connection error raises ParserError."""
        # Using a URL that will fail to connect (localhost with no server)
        with pytest.raises(ParserError):
            self.service.parse("http://localhost:59999/nonexistent")


class TestFilePathDetection:
    """Tests for file path detection."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
    
    def test_is_file_path_pdf(self):
        """Test PDF file path detection."""
        assert self.service._is_file_path("document.pdf") is True
        assert self.service._is_file_path("/path/to/document.PDF") is True
    
    def test_is_file_path_docx(self):
        """Test DOCX file path detection."""
        assert self.service._is_file_path("document.docx") is True
        assert self.service._is_file_path("/path/to/document.DOCX") is True
    
    def test_is_file_path_doc(self):
        """Test DOC file path detection."""
        assert self.service._is_file_path("document.doc") is True
    
    def test_is_file_path_txt(self):
        """Test TXT file path detection."""
        assert self.service._is_file_path("document.txt") is True
    
    def test_is_file_path_not_supported(self):
        """Test unsupported file extension detection."""
        assert self.service._is_file_path("document.html") is False
        assert self.service._is_file_path("document.jpg") is False


class TestFileParsing:
    """Tests for file parsing (TXT files)."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
        self.test_dir = Path(__file__).parent / "test_data"
        self.test_dir.mkdir(exist_ok=True)
    
    def teardown_method(self):
        """Clean up test files."""
        for f in self.test_dir.glob("*"):
            f.unlink()
        self.test_dir.rmdir()
    
    def test_parse_txt_file(self):
        """Test parsing a TXT file."""
        txt_file = self.test_dir / "test.txt"
        content = "This is test content.\nLine 2."
        txt_file.write_text(content)
        
        result = self.service.parse(str(txt_file))
        assert result == content.strip()
    
    def test_parse_nonexistent_file_raises_error(self):
        """Test that non-existent file raises ParserError."""
        with pytest.raises(ParserError) as exc_info:
            self.service.parse("nonexistent.pdf")
        assert "File not found" in str(exc_info.value)
    
    def test_parse_unsupported_extension_treated_as_text(self):
        """Test that paths with unsupported extensions are treated as plain text strings."""
        html_file = self.test_dir / "test.html"
        content = "<html>test</html>"
        html_file.write_text(content)
        
        # Paths with unsupported extensions are treated as plain text (the path string itself)
        # This is expected behavior - only supported extensions trigger file reading
        result = self.service.parse(str(html_file))
        # The result is the path string, not the file content
        assert "test.html" in result


class TestPDFParsing:
    """Tests for PDF parsing."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
        self.test_dir = Path(__file__).parent / "test_data"
        self.test_dir.mkdir(exist_ok=True)
    
    def teardown_method(self):
        """Clean up test files."""
        for f in self.test_dir.glob("*"):
            f.unlink()
        self.test_dir.rmdir()
    
    def test_parse_pdf_file(self):
        """Test parsing a PDF file."""
        import PyPDF2
        from PyPDF2 import PdfWriter
        
        # Create a simple PDF for testing
        pdf_file = self.test_dir / "test.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)  # Letter size
        
        with open(pdf_file, 'wb') as f:
            writer.write(f)
        
        # Parse the PDF (will return empty string since no text)
        result = self.service.parse(str(pdf_file))
        assert isinstance(result, str)
    
    def test_parse_nonexistent_pdf_raises_error(self):
        """Test that non-existent PDF raises ParserError."""
        with pytest.raises(ParserError) as exc_info:
            self.service.parse("nonexistent.pdf")
        assert "File not found" in str(exc_info.value)


class TestDOCXParsing:
    """Tests for DOCX parsing."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.service = ParserService()
        self.test_dir = Path(__file__).parent / "test_data"
        self.test_dir.mkdir(exist_ok=True)
    
    def teardown_method(self):
        """Clean up test files."""
        for f in self.test_dir.glob("*"):
            f.unlink()
        self.test_dir.rmdir()
    
    def test_parse_docx_file(self):
        """Test parsing a DOCX file."""
        from docx import Document
        
        # Create a simple DOCX for testing
        docx_file = self.test_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("This is a test paragraph.")
        doc.add_paragraph("This is another paragraph.")
        doc.save(str(docx_file))
        
        # Parse the DOCX
        result = self.service.parse(str(docx_file))
        assert "This is a test paragraph." in result
        assert "This is another paragraph." in result
    
    def test_parse_nonexistent_docx_raises_error(self):
        """Test that non-existent DOCX raises ParserError."""
        with pytest.raises(ParserError) as exc_info:
            self.service.parse("nonexistent.docx")
        assert "File not found" in str(exc_info.value)